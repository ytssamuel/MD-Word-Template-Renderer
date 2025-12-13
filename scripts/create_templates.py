#!/usr/bin/env python
"""
建立 Word 模板範例

使用 python-docx 建立多個模板範例
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_simple_template():
    """建立簡單模板"""
    doc = Document()
    
    # 標題
    title = doc.add_heading('變更單資料', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本資訊
    doc.add_heading('基本資訊', level=1)
    doc.add_paragraph('系統名稱：{{系統名稱}}')
    doc.add_paragraph('變更單號：{{變更單號}}')
    doc.add_paragraph('預定換版日期：{{預定換版日期}}')
    doc.add_paragraph('預定變更時段：{{預定變更時段}}')
    
    # 儲存
    output_path = Path('templates/simple_template.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✓ 建立 {output_path}")


def create_example_template():
    """建立範例模板（中等複雜度）"""
    doc = Document()
    
    # 標題
    title = doc.add_heading('變更單資料模板', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本資訊區塊
    doc.add_heading('【基本資訊】', level=1)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    # 注意：欄位名稱中有特殊字元需要用引號包起來
    fields = [
        ('系統名稱', '{{系統名稱}}'),
        ('預定換版日期', '{{預定換版日期}}'),
        ('預定變更時段', '{{預定變更時段}}'),
        ('變更單號', '{{變更單號}}'),
        ('備援路徑', '{{備援路徑}}'),
        ('需求依據', '{{data["需求依據(INC/PBI)"]}}'),
        ('資料庫', '{{資料庫}}'),
    ]
    
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    # 相關系統 - 使用字典存取法處理特殊字元
    doc.add_heading('【相關系統】', level=1)
    doc.add_paragraph('{% if 中介軟體 %}中介軟體：{{中介軟體}}{% endif %}')
    doc.add_paragraph('{% if data["網路(含防火牆)"] %}網路設定：{{data["網路(含防火牆)"]}}{% endif %}')
    doc.add_paragraph('{% if 其他系統 %}其他系統：{{其他系統}}{% endif %}')
    
    # 通知資訊
    doc.add_heading('【通知資訊】', level=1)
    doc.add_paragraph('計畫性維護：{{計畫性維護}}')
    doc.add_paragraph('通知對象：{{通知對象}}')
    doc.add_paragraph('通知方式：{{data["通知/公告方式"]}}')
    doc.add_paragraph('通知內容：{{data["通知/公告內容"]}}')
    
    # 異動說明
    doc.add_heading('【異動說明】', level=1)
    doc.add_paragraph('通報單異動原因：{{data["通報單-異動原因"]}}')
    
    # 測試案例（迴圈）
    doc.add_heading('【測試案例】', level=1)
    
    # 使用 Jinja2 迴圈語法 - 注意特殊字元欄位
    loop_para = doc.add_paragraph()
    loop_para.add_run('{% for item in data["異動內容-測試案例"] %}')
    
    doc.add_paragraph('{{loop.index}}. {{item.value}}')
    
    # 子項目迴圈
    child_para = doc.add_paragraph()
    child_para.add_run('{% for child in item.children %}')
    doc.add_paragraph('   {{child.number}}. {{child.value}}')
    child_end = doc.add_paragraph()
    child_end.add_run('{% endfor %}')
    
    loop_end = doc.add_paragraph()
    loop_end.add_run('{% endfor %}')
    
    # 測試環境
    doc.add_heading('【測試環境】', level=1)
    doc.add_paragraph('測試日期：{{測試日期}}')
    doc.add_paragraph('主機名稱：{{data["主機名稱(IP)"]}}')
    doc.add_paragraph('作業系統：{{作業系統}}')
    doc.add_paragraph('瀏覽器：{{瀏覽器}}')
    doc.add_paragraph('測試網址：{{測試網址}}')
    doc.add_paragraph('測試參數：{{測試參數}}')
    
    # 儲存
    output_path = Path('templates/example_template.docx')
    doc.save(output_path)
    print(f"✓ 建立 {output_path}")


def create_full_template():
    """建立完整功能模板"""
    doc = Document()
    
    # 標題
    title = doc.add_heading('變更作業確認單', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本資訊
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    version_para.add_run('文件編號：{{變更單號}}')
    
    doc.add_paragraph()  # 空行
    
    # ===== 第一部分：基本資訊 =====
    doc.add_heading('一、變更基本資訊', level=1)
    
    # 建立表格
    table1 = doc.add_table(rows=4, cols=4)
    table1.style = 'Table Grid'
    
    # 設定表格內容 - 使用字典存取法處理特殊字元
    cells_data = [
        [('系統名稱', 1), ('{{系統名稱}}', 3)],
        [('變更單號', 1), ('{{變更單號}}', 1), ('需求依據', 1), ('{{data["需求依據(INC/PBI)"]}}', 1)],
        [('預定換版日期', 1), ('{{預定換版日期}}', 1), ('變更時段', 1), ('{{預定變更時段}}', 1)],
        [('備援路徑', 1), ('{{備援路徑}}', 3)],
    ]
    
    for row_idx, row_data in enumerate(cells_data):
        col_idx = 0
        for text, span in row_data:
            cell = table1.rows[row_idx].cells[col_idx]
            cell.text = text
            col_idx += span
    
    doc.add_paragraph()  # 空行
    
    # ===== 第二部分：相關系統 =====
    doc.add_heading('二、相關系統設定', level=1)
    
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'
    
    # 使用字典存取法處理含括號的欄位名稱
    system_fields = [
        ('資料庫', '{{資料庫}}'),
        ('中介軟體', '{% if 中介軟體 %}{{中介軟體}}{% else %}無{% endif %}'),
        ('網路(含防火牆)', '{% if data["網路(含防火牆)"] %}{{data["網路(含防火牆)"]}}{% else %}無{% endif %}'),
        ('其他系統', '{% if 其他系統 %}{{其他系統}}{% else %}無{% endif %}'),
    ]
    
    for i, (label, value) in enumerate(system_fields):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # ===== 第三部分：通知資訊 =====
    doc.add_heading('三、通知相關資訊', level=1)
    
    table3 = doc.add_table(rows=4, cols=2)
    table3.style = 'Table Grid'
    
    # 使用字典存取法處理含斜線的欄位名稱
    notify_fields = [
        ('計畫性維護', '{{計畫性維護}}'),
        ('通知對象', '{{通知對象}}'),
        ('通知方式', '{{data["通知/公告方式"]}}'),
        ('通知內容', '{{data["通知/公告內容"]}}'),
    ]
    
    for i, (label, value) in enumerate(notify_fields):
        table3.rows[i].cells[0].text = label
        table3.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # ===== 第四部分：異動原因 =====
    doc.add_heading('四、異動原因說明', level=1)
    doc.add_paragraph('{{data["通報單-異動原因"]}}')
    
    doc.add_paragraph()
    
    # ===== 第五部分：測試案例 =====
    doc.add_heading('五、測試案例', level=1)
    
    # 使用普通 for 迴圈（非 {%tr} 表格列迴圈）
    doc.add_paragraph('{% for item in data["異動內容-測試案例"] %}')
    
    # 每個測試案例用獨立段落顯示
    doc.add_paragraph('【案例 {{loop.index}}】{{item.value}}')
    doc.add_paragraph('  子項目：{% for c in item.children %}{{c.number}}.{{c.value}} {% endfor %}')
    
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph()
    
    # ===== 第六部分：測試環境 =====
    doc.add_heading('六、測試環境資訊', level=1)
    
    table4 = doc.add_table(rows=6, cols=2)
    table4.style = 'Table Grid'
    
    # 使用字典存取法處理含括號的欄位名稱
    env_fields = [
        ('測試日期', '{{測試日期}}'),
        ('主機名稱(IP)', '{{data["主機名稱(IP)"]}}'),
        ('作業系統', '{{作業系統}}'),
        ('瀏覽器', '{{瀏覽器}}'),
        ('測試網址', '{{測試網址}}'),
        ('測試參數', '{{測試參數}}'),
    ]
    
    for i, (label, value) in enumerate(env_fields):
        table4.rows[i].cells[0].text = label
        table4.rows[i].cells[1].text = value
    
    # 頁尾
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('--- 文件結束 ---')
    
    # 儲存
    output_path = Path('templates/full_template.docx')
    doc.save(output_path)
    print(f"✓ 建立 {output_path}")


def main():
    """建立所有模板"""
    print("開始建立 Word 模板範例...\n")
    
    create_simple_template()
    create_example_template()
    create_full_template()
    
    print("\n✓ 所有模板建立完成！")
    print("\n模板位置：")
    print("  - templates/simple_template.docx  (簡單模板)")
    print("  - templates/example_template.docx (中等複雜度)")
    print("  - templates/full_template.docx    (完整功能)")


if __name__ == '__main__':
    main()
